
from fastapi import APIRouter, Request
from typing import List, Optional
from fastapi.responses import JSONResponse
from typing import Optional

from controller.query_processor import query_processor, query_processor2
from fastapi import FastAPI, File, UploadFile,Form
from PIL import Image
import io
import speech_recognition as sr
# import moviepy.editor as mp
import PyPDF2
from docx import Document
router = APIRouter()


@router.get("/response")
async def response(query: Optional[str] = None,file: Optional[UploadFile] = None):
    filename = file.filename.lower() if file else None
    content = await file.read() if file else None
    result = await query_processor(query=query)




    # return JSONResponse(content=result)
    return result




    content = await file.read() if file else None
    image = None
    audio = None
    video = None
    file_text = None
    result = None

    if filename:
        if filename.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
            try:
                image = Image.open(io.BytesIO(content))
            except Exception as e:
                result = f"Image read error: {e}"

        elif filename.endswith(('.wav', '.mp3', '.m4a')):
            try:
                recognizer = sr.Recognizer()
                with sr.AudioFile(io.BytesIO(content)) as source:
                    audio = recognizer.record(source)
                # Optionally transcribe
                try:
                    file_text = recognizer.recognize_google(audio)
                except Exception:
                    file_text = None
            except Exception as e:
                result = f"Audio read error: {e}"

        elif filename.endswith(('.mp4', '.avi', '.mov', '.mkv')):
            try:
                # video = mp.VideoFileClip(io.BytesIO(content))
                video = content  # Placeholder: pass 
            except Exception as e:
                result = f"Video read error: {e}"

        elif filename.endswith('.pdf'):
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                file_text = ""
                for page in pdf_reader.pages:
                    file_text += page.extract_text()
            except Exception as e:
                result = f"PDF read error: {e}"

        elif filename.endswith('.docx'):
            try:
                doc = Document(io.BytesIO(content))
                file_text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                result = f"DOCX read error: {e}"
        else:
            result = "Unsupported file type."

    # If text is provided, use it as file_text
    if text:
        query = text
        result = "Text successfully read."

    # Send all extracted content to query_processor
    llm_result = await query_processor(
        query=query,
        file_text=file_text,
        image=image,
        audio=audio,
        video=video,
    )
    return llm_result

@router.post("/ask")
async def ask(
    query: Optional[str] = Form(None),
    model: Optional[str] = Form("gemini-2.5-flash"),
    files: Optional[List[UploadFile]] = File(None),
):
    
    if files:
        recognizer = sr.Recognizer()
        for f in files:
            try:
                filename = (f.filename or "").lower()
                raw = await f.read()

                if filename.endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp")):
                    # keep raw image for multimodal
                    uploads.append({"bytes": raw, "mime": "image/*"})

                elif filename.endswith((".wav", ".mp3", ".m4a", ".flac", ".ogg")):
                    # keep raw audio, try speech-to-text best-effort
                    uploads.append({"bytes": raw, "mime": "audio/*"})
                    try:
                        with sr.AudioFile(io.BytesIO(raw)) as source:
                            audio_data = recognizer.record(source)
                        stt_text = recognizer.recognize_google(audio_data)
                        if stt_text:
                            extracted_texts.append(stt_text)
                    except Exception:
                        pass  # non-fatal

                elif filename.endswith((".mp4", ".avi", ".mov", ".mkv", ".webm")):
                    uploads.append({"bytes": raw, "mime": "video/*"})

                elif filename.endswith(".pdf"):
                    try:
                        reader = PyPDF2.PdfReader(io.BytesIO(raw))
                        pdf_text = ""
                        for page in reader.pages:
                            pdf_text += page.extract_text() or ""
                        if pdf_text.strip():
                            extracted_texts.append(pdf_text)
                    except Exception as e:
                        return JSONResponse({"error": f"PDF read error: {e}"}, status_code=400)

                elif filename.endswith(".docx"):
                    try:
                        doc = Document(io.BytesIO(raw))
                        doc_text = "\n".join(p.text for p in doc.paragraphs)
                        if doc_text.strip():
                            extracted_texts.append(doc_text)
                    except Exception as e:
                        return JSONResponse({"error": f"DOCX read error: {e}"}, status_code=400)

                else:
                    # Unknown file type—still pass bytes along as generic binary
                    uploads.append({"bytes": raw, "mime": "application/octet-stream"})

            except Exception as e:
                return JSONResponse({"error": f"File processing error: {e}"}, status_code=400)

    # Join all extracted text into one block; let the processor fuse with RAG context
    extracted_text = "\n\n".join(t for t in extracted_texts if t.strip())

    # Delegate to the processor (Gemini 2.5 Flash by default)
    result = await query_processor2(
        query=query or "",
        model=model,
        uploads=uploads,
        extracted_text=extracted_text,
    )
    return result